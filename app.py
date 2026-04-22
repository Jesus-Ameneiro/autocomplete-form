"""
Template Document Filler — Streamlit App (v3)
Document View with live inline preview.
- Text placeholders are global (shared across all occurrences).
- Numeric/currency placeholders in tables are independent, named by Header_Row.
Exports to DOCX and PDF.
"""

import re, io, os, html as html_mod
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

# ─── Constants ───────────────────────────────────────────────────────────────

PLACEHOLDER_RE = re.compile(r"\[([^\[\]]+)\]")
WD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
CURRENCY_CHARS = set("$€£¥")

MULTILINE_KEYWORDS = (
    "descripción", "description", "notas", "notes", "dirección", "address",
    "párrafo", "paragraph", "cuerpo", "body", "contenido", "content",
    "detalles", "details", "resumen", "summary", "comentarios", "comments",
)


# ─── Classify placeholder content ───────────────────────────────────────────

def _classify_placeholder(raw_text):
    """
    Classify a placeholder's inner text.
    Returns (type, prefix, label) where:
      type:   'currency' | 'numeric' | 'text'
      prefix: currency symbol (e.g. '$') or ''
      label:  text after the prefix (stripped)
    """
    s = raw_text.strip()
    if not s:
        return "text", "", s
    if s[0] in CURRENCY_CHARS:
        return "currency", s[0], s[1:].strip()
    if re.match(r"^[\d,.\s]+$", s):
        return "numeric", "", s
    return "text", "", s


def _header_to_field(header_text):
    """Convert column header to a clean field name segment: 'Precio Unitario' → 'Precio_Unitario'."""
    return re.sub(r"\s+", "_", header_text.strip())


# ─── Document Parsing Helpers ────────────────────────────────────────────────

def _is_list_paragraph(para):
    ppr = para._element.find(f".//{{{WD_NS}}}pPr")
    if ppr is not None:
        return ppr.find(f".//{{{WD_NS}}}numPr") is not None
    return False


def _get_alignment(para):
    a = para.alignment
    if a == WD_ALIGN_PARAGRAPH.RIGHT:
        return "right"
    elif a == WD_ALIGN_PARAGRAPH.CENTER:
        return "center"
    elif a == WD_ALIGN_PARAGRAPH.JUSTIFY:
        return "justify"
    return "left"


# ─── Iterate every paragraph (for global placeholder scanning) ──────────────

def _iter_body_paragraphs(doc):
    """Yield paragraphs from body only (NOT tables)."""
    yield from doc.paragraphs


def _iter_header_footer_paragraphs(doc):
    """Yield paragraphs from headers/footers."""
    for section in doc.sections:
        for hf in (
            section.header, section.footer,
            section.first_page_header, section.first_page_footer,
            section.even_page_header, section.even_page_footer,
        ):
            if hf is not None and hf.is_linked_to_previous is False:
                yield from hf.paragraphs
                for table in hf.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            yield from cell.paragraphs


def iter_all_paragraphs(doc):
    """Yield every paragraph from body, tables, headers, footers."""
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
                for nested in cell.tables:
                    for nr in nested.rows:
                        for nc in nr.cells:
                            yield from nc.paragraphs
    yield from _iter_header_footer_paragraphs(doc)


# ─── Placeholder Discovery ──────────────────────────────────────────────────

def find_placeholders(doc):
    """
    Scan the document and return:
      global_phs:    ordered list of unique global (text) placeholder names
      global_lower:  lowercase → display name map for globals
      table_fields:  list of dicts for each numeric/currency cell in tables:
                     {name, prefix, table_idx, row_idx, col_idx, original}
      table_text_cells: list of (table_idx, row_idx, col_idx, placeholder_name)
                        for text placeholders in tables (use global values)
    """
    # --- Pass 1: body + header/footer paragraphs → global text placeholders ---
    seen_lower = {}
    order = []

    def _add_global(name):
        key = name.lower()
        if key not in seen_lower:
            seen_lower[key] = name
            order.append(name)

    for para in _iter_body_paragraphs(doc):
        full_text = "".join(r.text for r in para.runs)
        for m in PLACEHOLDER_RE.finditer(full_text):
            raw = m.group(1).strip()
            ptype, prefix, label = _classify_placeholder(raw)
            if ptype == "text":
                _add_global(raw)

    for para in _iter_header_footer_paragraphs(doc):
        full_text = "".join(r.text for r in para.runs)
        for m in PLACEHOLDER_RE.finditer(full_text):
            raw = m.group(1).strip()
            ptype, _, _ = _classify_placeholder(raw)
            if ptype == "text":
                _add_global(raw)

    # --- Pass 2: tables → classify each cell's placeholder ---
    table_fields = []
    table_text_cells = []

    for ti, table in enumerate(doc.tables):
        if len(table.rows) < 2:
            continue
        headers = [c.text.strip() for c in table.rows[0].cells]

        for ri in range(1, len(table.rows)):
            row = table.rows[ri]
            data_row_num = ri  # 1-indexed data row
            for ci, cell in enumerate(row.cells):
                full_text = "".join(
                    r.text for p in cell.paragraphs for r in p.runs
                )
                match = PLACEHOLDER_RE.search(full_text)
                if not match:
                    continue

                raw = match.group(1).strip()
                ptype, prefix, label = _classify_placeholder(raw)

                if ptype in ("currency", "numeric"):
                    header_label = _header_to_field(headers[ci]) if ci < len(headers) else f"Col{ci}"
                    field_name = f"{header_label}_{data_row_num}"
                    table_fields.append({
                        "name": field_name,
                        "prefix": prefix,
                        "table_idx": ti,
                        "row_idx": ri,
                        "col_idx": ci,
                        "original": raw,
                    })
                else:
                    # Text placeholder in table → global
                    _add_global(raw)
                    table_text_cells.append((ti, ri, ci, raw))

    return order, seen_lower, table_fields, table_text_cells


# ─── Run-aware replacement helpers ──────────────────────────────────────────

def _replace_in_paragraph(paragraph, repl_lower):
    """Replace [placeholder] tokens in a paragraph, preserving run formatting."""
    runs = paragraph.runs
    if not runs:
        return
    char_origins = []
    for ri, run in enumerate(runs):
        for ci in range(len(run.text)):
            char_origins.append((ri, ci))
    full_text = "".join(r.text for r in runs)
    matches = []
    for m in PLACEHOLDER_RE.finditer(full_text):
        key = m.group(1).strip().lower()
        if key in repl_lower:
            matches.append((m.start(), m.end(), repl_lower[key]))
    if not matches:
        return
    for start, end, value in reversed(matches):
        first_ri, first_off = char_origins[start]
        last_ri, last_off = char_origins[end - 1]
        if first_ri == last_ri:
            r = runs[first_ri]
            r.text = r.text[:first_off] + value + r.text[last_off + 1:]
        else:
            runs[first_ri].text = runs[first_ri].text[:first_off] + value
            runs[last_ri].text = runs[last_ri].text[last_off + 1:]
            for mid in range(first_ri + 1, last_ri):
                runs[mid].text = ""


def _replace_cell(cell, value):
    """Replace the entire placeholder in a table cell with a value."""
    for para in cell.paragraphs:
        runs = para.runs
        if not runs:
            continue
        full_text = "".join(r.text for r in runs)
        if not PLACEHOLDER_RE.search(full_text):
            continue
        new_text = PLACEHOLDER_RE.sub(value, full_text)
        runs[0].text = new_text
        for r in runs[1:]:
            r.text = ""


def apply_replacements(doc, global_values, table_field_values, table_fields):
    """
    Replace placeholders in the document:
    - Body/header/footer: use global_values (case-insensitive)
    - Table numeric/currency cells: use table_field_values with prefix
    """
    repl_lower = {k.lower(): v for k, v in global_values.items()}

    # Body + header/footer paragraphs (global)
    for para in _iter_body_paragraphs(doc):
        _replace_in_paragraph(para, repl_lower)
    for para in _iter_header_footer_paragraphs(doc):
        _replace_in_paragraph(para, repl_lower)

    # Table text cells (global)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, repl_lower)

    # Table numeric/currency cells (independent, with prefix)
    for tf in table_fields:
        val = table_field_values.get(tf["name"], "")
        if not val.strip():
            continue
        display_val = tf["prefix"] + val
        table = doc.tables[tf["table_idx"]]
        cell = table.rows[tf["row_idx"]].cells[tf["col_idx"]]
        _replace_cell(cell, display_val)


# ─── HTML Rendering ─────────────────────────────────────────────────────────

def _apply_run_fmt(html_fragment, run):
    if not run:
        return html_fragment
    is_bold = getattr(run, "bold", False)
    is_italic = getattr(run, "italic", False)
    font_size = None
    if run.font and run.font.size:
        font_size = round(run.font.size / 12700)
    if is_bold:
        html_fragment = f"<strong>{html_fragment}</strong>"
    if is_italic:
        html_fragment = f"<em>{html_fragment}</em>"
    if font_size and font_size != 12:
        html_fragment = f'<span style="font-size:{font_size}pt">{html_fragment}</span>'
    return html_fragment


def render_runs_html(runs, mode, global_values, global_lower):
    """Render runs to HTML for body/header/footer paragraphs (global placeholders)."""
    if not runs:
        return ""
    full_text = ""
    char_run = []
    for ri, run in enumerate(runs):
        for _ in run.text:
            char_run.append(ri)
        full_text += run.text
    if not full_text.strip():
        return ""

    ph_spans = [
        (m.start(), m.end(), m.group(1).strip())
        for m in PLACEHOLDER_RE.finditer(full_text)
    ]

    parts = []
    pos = 0

    def _emit_plain(start, end):
        if start >= end:
            return
        i = start
        while i < end:
            ri = char_run[i]
            j = i + 1
            while j < end and char_run[j] == ri:
                j += 1
            fragment = html_mod.escape(full_text[i:j])
            parts.append(_apply_run_fmt(fragment, runs[ri]))
            i = j

    for ph_start, ph_end, ph_name in ph_spans:
        _emit_plain(pos, ph_start)
        key = ph_name.lower()
        display_name = global_lower.get(key, ph_name)
        val = global_values.get(display_name, "")
        fmt_run = runs[char_run[ph_start]]

        if mode == "preview":
            inner = html_mod.escape(val) if val.strip() else html_mod.escape(full_text[ph_start:ph_end])
        else:  # edit
            if val.strip():
                inner = f'<span class="filled-value">{html_mod.escape(val)}</span>'
            else:
                inner = f'<span class="empty-placeholder">{html_mod.escape(ph_name)}</span>'

        parts.append(_apply_run_fmt(inner, fmt_run))
        pos = ph_end

    _emit_plain(pos, len(full_text))
    return "".join(parts)


def _render_table_cell_html(cell, cell_key, mode, global_values, global_lower, tf_map, table_field_values):
    """
    Render a table cell. If it has a table-field mapping (numeric/currency),
    use the independent field value. Otherwise, use global replacements.
    """
    tf = tf_map.get(cell_key)  # cell_key is (table_element_id, row_idx, col_idx)

    if tf is not None:
        # This is a numeric/currency cell — use independent field value
        val = table_field_values.get(tf["name"], "")
        prefix = tf["prefix"]
        field_name = tf["name"]

        if mode == "preview":
            if val.strip():
                return html_mod.escape(prefix + val)
            else:
                return html_mod.escape(prefix + "___")
        else:  # edit
            if val.strip():
                display = prefix + val
                return f'<span class="filled-value">{html_mod.escape(display)}</span>'
            else:
                return f'<span class="empty-placeholder">{html_mod.escape(field_name)}</span>'
    else:
        # Normal cell — use global placeholder rendering
        cell_parts = []
        for para in cell.paragraphs:
            cell_parts.append(render_runs_html(para.runs, mode, global_values, global_lower))
        return " ".join(p for p in cell_parts if p) or "&nbsp;"


# ─── Parse Document Structure ───────────────────────────────────────────────

def parse_document(doc):
    elements = []
    body = doc.element.body
    para_map = {id(p._element): p for p in doc.paragraphs}
    table_map = {id(t._element): t for t in doc.tables}
    current_list_items = []

    for child in body:
        tag = etree.QName(child.tag).localname
        if tag == "p":
            para = para_map.get(id(child))
            if para is None:
                continue
            if _is_list_paragraph(para):
                current_list_items.append(para)
            else:
                if current_list_items:
                    elements.append({"type": "list", "items": current_list_items})
                    current_list_items = []
                elements.append({"type": "paragraph", "para": para})
        elif tag == "tbl":
            if current_list_items:
                elements.append({"type": "list", "items": current_list_items})
                current_list_items = []
            tbl = table_map.get(id(child))
            if tbl:
                elements.append({"type": "table", "table": tbl})

    if current_list_items:
        elements.append({"type": "list", "items": current_list_items})
    return elements


# ─── Full Document HTML Rendering ───────────────────────────────────────────

def render_document_html(doc, mode, global_values, global_lower, table_fields, table_field_values):
    """Render the full document as styled HTML."""

    # Build (table_element_id, row_idx, col_idx) → table_field mapping
    tf_map = {}
    for tf in table_fields:
        table = doc.tables[tf["table_idx"]]
        key = (id(table._tbl), tf["row_idx"], tf["col_idx"])
        tf_map[key] = tf

    elements = parse_document(doc)
    parts = []

    for el in elements:
        if el["type"] == "paragraph":
            para = el["para"]
            align = _get_alignment(para)
            runs = para.runs
            if not runs:
                parts.append(f'<p style="text-align:{align};margin:4px 0">&nbsp;</p>')
                continue
            inner = render_runs_html(runs, mode, global_values, global_lower)
            if not inner.strip():
                inner = "&nbsp;"
            parts.append(f'<p style="text-align:{align};margin:6px 0;line-height:1.5">{inner}</p>')

        elif el["type"] == "table":
            table = el["table"]
            tbl_id = id(table._tbl)
            rows_html = []
            for ri, row in enumerate(table.rows):
                cells_html = []
                for ci, cell in enumerate(row.cells):
                    cell_key = (tbl_id, ri, ci)
                    inner = _render_table_cell_html(
                        cell, cell_key, mode, global_values, global_lower,
                        tf_map, table_field_values,
                    )
                    tag = "th" if ri == 0 else "td"
                    cells_html.append(f"<{tag}>{inner}</{tag}>")
                rows_html.append("<tr>" + "".join(cells_html) + "</tr>")
            parts.append('<table class="doc-table">' + "".join(rows_html) + "</table>")

        elif el["type"] == "list":
            lis = []
            for para in el["items"]:
                inner = render_runs_html(para.runs, mode, global_values, global_lower)
                lis.append(f"<li>{inner}</li>")
            parts.append("<ul>" + "".join(lis) + "</ul>")

    body_html = "\n".join(parts)
    css = _get_display_css()
    return f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><style>{css}</style></head>
<body><div class="doc-page">{body_html}</div></body></html>"""


def _get_display_css():
    return """
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700&display=swap');
    * { box-sizing: border-box; margin: 0; padding: 0; }
    body {
        font-family: 'Inter', 'Segoe UI', Arial, sans-serif;
        font-size: 11pt; color: #1a1a1a; background: transparent;
    }
    .doc-page {
        max-width: 720px; margin: 0 auto; padding: 32px 40px;
        background: white; border: 1px solid #d0d0d0; border-radius: 4px;
        box-shadow: 0 1px 4px rgba(0,0,0,0.06); min-height: 600px;
    }
    p { margin: 6px 0; line-height: 1.55; }
    ul { margin: 8px 0 8px 28px; }
    li { margin: 4px 0; line-height: 1.55; }
    .doc-table {
        width: 100%; border-collapse: collapse; margin: 12px 0; font-size: 10pt;
    }
    .doc-table th, .doc-table td {
        border: 1px solid #999; padding: 6px 10px; text-align: left;
    }
    .doc-table th { background: #f0f0f0; font-weight: 700; }
    .filled-value {
        background: #d4edda; color: #155724;
        padding: 1px 5px; border-radius: 3px;
        border-bottom: 2px solid #28a745; font-weight: 600;
    }
    .empty-placeholder {
        background: #fff3cd; color: #856404;
        padding: 1px 5px; border-radius: 3px;
        border: 1px dashed #ffc107; font-style: italic; font-size: 0.92em;
    }
    """


# ─── PDF Generation via fpdf2 ───────────────────────────────────────────────

def generate_pdf(doc, global_values, global_lower, table_fields, table_field_values):
    from fpdf import FPDF

    pdf = FPDF(orientation="P", unit="mm", format="Letter")
    pdf.set_auto_page_break(auto=True, margin=20)

    # Try to find a Unicode TTF font; search several common paths
    fn = "Helvetica"  # safe fallback (ASCII only)
    search_dirs = [
        "/usr/share/fonts/truetype/dejavu",
        "/usr/share/fonts/truetype/liberation",
        "/usr/share/fonts/TTF",
        "/usr/local/share/fonts",
    ]
    font_variants = {
        "dejavu": {
            "": "DejaVuSans.ttf",
            "B": "DejaVuSans-Bold.ttf",
            "I": "DejaVuSans-Oblique.ttf",
            "BI": "DejaVuSans-BoldOblique.ttf",
        },
        "liberation": {
            "": "LiberationSans-Regular.ttf",
            "B": "LiberationSans-Bold.ttf",
            "I": "LiberationSans-Italic.ttf",
            "BI": "LiberationSans-BoldItalic.ttf",
        },
    }

    font_loaded = False
    for fdir in search_dirs:
        if not os.path.isdir(fdir):
            continue
        for family_name, variants in font_variants.items():
            regular = os.path.join(fdir, variants[""])
            if not os.path.isfile(regular):
                continue
            try:
                fn = family_name
                pdf.add_font(fn, "", regular)
                # Add bold/italic only if the files actually exist
                for style_key in ("B", "I", "BI"):
                    path = os.path.join(fdir, variants[style_key])
                    if os.path.isfile(path):
                        pdf.add_font(fn, style_key, path)
                font_loaded = True
                break
            except Exception:
                fn = "Helvetica"
                continue
        if font_loaded:
            break
    else:
        fn = "Helvetica"

    bullet_char = "\u2022" if font_loaded else "-"

    pdf.add_page()
    pdf.set_margins(18, 18, 18)
    pdf.set_font(fn, size=11)

    # Build global replacement dict
    repl_lower = {k.lower(): v for k, v in global_values.items()}

    # Build table field lookup: (table_idx, row_idx, col_idx) → (prefix, value)
    tf_lookup = {}
    for tf in table_fields:
        val = table_field_values.get(tf["name"], "")
        tf_lookup[(tf["table_idx"], tf["row_idx"], tf["col_idx"])] = (tf["prefix"], val)

    def _resolve_global(text):
        def _sub(m):
            key = m.group(1).strip().lower()
            display = global_lower.get(key, m.group(1).strip())
            val = global_values.get(display, "")
            return val if val.strip() else m.group(0)
        return PLACEHOLDER_RE.sub(_sub, text)

    elements = parse_document(doc)
    table_counter = -1

    for el in elements:
        if el["type"] == "paragraph":
            para = el["para"]
            runs = para.runs
            full_text = "".join(r.text for r in runs)
            resolved = _resolve_global(full_text)
            if not resolved.strip():
                pdf.ln(4)
                continue
            align_map = {"left": "L", "right": "R", "center": "C", "justify": "J"}
            pdf_align = align_map.get(_get_alignment(para), "L")
            font_size = 11
            if runs and runs[0].font and runs[0].font.size:
                font_size = round(runs[0].font.size / 12700)
            is_bold = all(r.bold for r in runs if r.text.strip()) if runs else False
            pdf.set_font(fn, "B" if is_bold else "", font_size)
            pdf.multi_cell(0, font_size * 0.45, resolved, align=pdf_align)
            pdf.ln(1)

        elif el["type"] == "table":
            table_counter += 1
            table = el["table"]
            num_cols = len(table.rows[0].cells) if table.rows else 0
            if num_cols == 0:
                continue
            col_w = (pdf.w - pdf.l_margin - pdf.r_margin) / num_cols

            for ri, row in enumerate(table.rows):
                is_hdr = ri == 0
                pdf.set_font(fn, "B" if is_hdr else "", 9)
                rh = 7
                for ci, cell in enumerate(row.cells):
                    key = (table_counter, ri, ci)
                    if key in tf_lookup:
                        prefix, val = tf_lookup[key]
                        cell_text = (prefix + val) if val.strip() else cell.text
                    else:
                        cell_text = _resolve_global(cell.text)
                    x0, y0 = pdf.get_x(), pdf.get_y()
                    pdf.rect(x0, y0, col_w, rh)
                    if is_hdr:
                        pdf.set_fill_color(235, 235, 235)
                        pdf.rect(x0, y0, col_w, rh, "F")
                    pdf.set_xy(x0 + 1, y0 + 1)
                    pdf.cell(col_w - 2, rh - 2, cell_text, align="L")
                    pdf.set_xy(x0 + col_w, y0)
                pdf.ln(rh)
            pdf.ln(2)

        elif el["type"] == "list":
            for para in el["items"]:
                full_text = "".join(r.text for r in para.runs)
                resolved = _resolve_global(full_text)
                pdf.set_font(fn, "", 11)
                pdf.cell(8, 5, bullet_char, align="R")
                pdf.set_font(fn, "", 11)
                pdf.multi_cell(0, 5, " " + resolved, align="J")
                pdf.ln(1)

    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf.getvalue()


# ─── Streamlit App ───────────────────────────────────────────────────────────

st.set_page_config(page_title="Template Filler", page_icon="📝", layout="wide")

st.markdown("""
<style>
    .block-container { max-width: 1200px; padding-top: 1.5rem; }
    div[data-testid="stFileUploader"] { margin-bottom: 0.5rem; }
    .placeholder-tag {
        display: inline-block; background: #e8f0fe; color: #1a56db;
        padding: 2px 8px; border-radius: 4px; font-family: monospace;
        font-size: 0.82em; margin: 2px;
    }
</style>
""", unsafe_allow_html=True)

st.title("📝 Template Document Filler")
st.caption(
    "Upload a `.docx` template with `[placeholder]` fields. "
    "Fill in the values and export as DOCX or PDF."
)

uploaded_file = st.file_uploader("Upload your Word template", type=["docx"])

if uploaded_file is not None:
    file_bytes = uploaded_file.read()
    doc = Document(io.BytesIO(file_bytes))
    global_phs, global_lower, table_fields, table_text_cells = find_placeholders(doc)

    if not global_phs and not table_fields:
        st.warning("No `[placeholder]` fields found. Use square brackets, e.g. `[Name]`.")
        st.stop()

    # ── Session state ────────────────────────────────────────────────
    if "g_vals" not in st.session_state:
        st.session_state.g_vals = {ph: "" for ph in global_phs}
    else:
        for ph in global_phs:
            if ph not in st.session_state.g_vals:
                st.session_state.g_vals[ph] = ""

    if "t_vals" not in st.session_state:
        st.session_state.t_vals = {tf["name"]: "" for tf in table_fields}
    else:
        for tf in table_fields:
            if tf["name"] not in st.session_state.t_vals:
                st.session_state.t_vals[tf["name"]] = ""

    # ── Helper: render input fields ──────────────────────────────────
    def _render_inputs():
        """Render all input fields. Returns (global_vals_dict, table_vals_dict)."""
        g = {}
        t = {}

        if global_phs:
            st.markdown("**Text Fields**")
            for ph in global_phs:
                is_ml = any(kw in ph.lower() for kw in MULTILINE_KEYWORDS)
                if is_ml:
                    g[ph] = st.text_area(
                        ph, value=st.session_state.g_vals.get(ph, ""),
                        key=f"doc_g_{ph}", height=80,
                    )
                else:
                    g[ph] = st.text_input(
                        ph, value=st.session_state.g_vals.get(ph, ""),
                        key=f"doc_g_{ph}",
                    )

        if table_fields:
            st.markdown("**Numeric / Currency Fields** *(table)*")
            for tf in table_fields:
                prefix_label = f" ({tf['prefix']})" if tf["prefix"] else ""
                label = f"{tf['name']}{prefix_label}"
                t[tf["name"]] = st.text_input(
                    label,
                    value=st.session_state.t_vals.get(tf["name"], ""),
                    key=f"doc_t_{tf['name']}",
                    placeholder=f"{tf['prefix']}...",
                )

        return g, t

    # ── Layout ───────────────────────────────────────────────────────
    left_col, right_col = st.columns([1, 2], gap="large")

    with left_col:
        st.markdown("#### Fields")
        total = len(global_phs) + len(table_fields)
        st.caption(f"{total} field(s) detected")
        g_vals, t_vals = _render_inputs()
        st.session_state.g_vals.update(g_vals)
        st.session_state.t_vals.update(t_vals)

    with right_col:
        preview_on = st.toggle("Preview mode", value=False, key="pv_toggle",
                               help="Toggle to see the clean final version")
        mode = "preview" if preview_on else "edit"
        doc_html = render_document_html(
            doc, mode, st.session_state.g_vals, global_lower,
            table_fields, st.session_state.t_vals,
        )
        st.components.v1.html(doc_html, height=900, scrolling=True)

    # ── Downloads (generated on click, not eagerly) ──────────────────
    st.markdown("---")
    st.markdown("#### Download")
    c1, c2, _ = st.columns([1, 1, 2])

    with c1:
        fresh = Document(io.BytesIO(file_bytes))
        apply_replacements(
            fresh, st.session_state.g_vals, st.session_state.t_vals, table_fields
        )
        buf = io.BytesIO()
        fresh.save(buf)
        st.download_button(
            "⬇️ Download DOCX", data=buf.getvalue(),
            file_name=uploaded_file.name.replace(".docx", "_filled.docx"),
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

    with c2:
        if st.button("📄 Generate PDF", use_container_width=True):
            with st.spinner("Generating PDF..."):
                pdf_bytes = generate_pdf(
                    doc, st.session_state.g_vals, global_lower,
                    table_fields, st.session_state.t_vals,
                )
            st.download_button(
                "⬇️ Download PDF", data=pdf_bytes,
                file_name=uploaded_file.name.replace(".docx", "_filled.pdf"),
                mime="application/pdf",
                use_container_width=True, key="dl_pdf_ready",
            )

else:
    st.markdown("---")
    st.markdown("**How it works**")
    st.markdown(
        "1. Create a Word document with placeholders in square brackets, "
        "e.g. `[Client Name]`, `[Invoice Date]`.\n"
        "2. Numeric/currency values in tables (like `[$0.00]` or `[0]`) are "
        "treated as independent fields, named by their column header and row.\n"
        "3. Upload the template, fill in the fields, and preview live.\n"
        "4. Download as **DOCX** or **PDF**."
    )
